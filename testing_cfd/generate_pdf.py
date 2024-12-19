import argparse
from docx import Document
from faker import Faker
import re
import datetime as dt
from datetime import datetime, timedelta
import random
import sys
import os


sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'config')))
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'features')))

import constants
from app_logger import logger
from db_functions import connect_to_oracle, exe_sql_query

# Initialize Faker for generating fake data
faker = Faker()

DEFAULT_ORDER_NUMBER = "12345"
DEFAULT_PICKSLIP_NUMBER = DEFAULT_ORDER_NUMBER
# DEFAULT_ORDER_DATE = datetime.now().strftime("%m-%d-%Y")
DEFAULT_ORDER_DATE = "11-20-2024"

def get_random_number(n):
    """
    Get a random number between 0 and n.
    :param n: The upper limit for the random number.
    :return: A random number between 0 and n.
    """
    return random.randint(0, n)


def get_order_info_from_table(ora_curx):
    """
    Get order number and date from a database table.
    :return: A order number and order date.
    """
    sql_query = f"SELECT ORDER_NUMBER, ORDER_DATE FROM {constants.ORDER_TABLE} ORDER BY CREATION_DATE DESC FETCH FIRST 1 ROW ONLY" # Get the oldest order
    # logger.info(f"ora_curx.version: {ora_curx.execute("SELECT * from V$VERSION")}")
    result, response = exe_sql_query(ora_curx, sql_query)
    if not result or not response:
        return DEFAULT_ORDER_NUMBER, DEFAULT_ORDER_DATE
    else:
        # random_ord_num = get_random_number(len(response) - 1)
        random_ord_num = 0
        logger.info(f"random_ord_num: {random_ord_num} response: {response[random_ord_num]}")
        logger.info(f"random_ord_num][0]: {response[random_ord_num][0]} response: {response[random_ord_num][1]}")
        ord_num_str = str(response[random_ord_num][0]+1)
        if response[random_ord_num][1]:
            # date_str = response[random_ord_num][1].strftime("%d-%m-%Y")
            date_str = response[random_ord_num][1].strftime("%d-%b-%y")
        else:
            date_str = DEFAULT_ORDER_DATE
        logger.info (f"order number: {ord_num_str} order date: {date_str}")
        return ord_num_str, date_str
    
def get_ord_po_from_table(ora_curx, document_type, order_type):
    """
    Get order number and date from a database table.
    :return: A order number and order date.
    """
    tbl_name = f"{document_type}_TABLE"
    tname = constants.__dict__[tbl_name]
    if order_type == 'ORDER':
        sql_query = f"SELECT ORDER_NUMBER FROM {tname} ORDER BY CREATION_DATE DESC FETCH FIRST 1 ROW ONLY" # Get the latest order
        increment = 1
    else:
        sql_query = f"SELECT PO_NUMBER FROM {constants.PO_TABLE} ORDER BY CREATION_DATE DESC FETCH FIRST 1 ROW ONLY" # Get the latest PO
        increment = 0
    result, response = exe_sql_query(ora_curx, sql_query)
    if not result or not response:
        return DEFAULT_PICKSLIP_NUMBER
    else:
        num_str = str(response[0][0]+increment)
        logger.info (f"num_str: {num_str}")
        return num_str


def get_ord_date_from_table(ora_curx):
    """
    Get order number and date from a database table.
    :return: A order number and order date.
    """
    sql_query = f"SELECT ORDER_DATE FROM {constants.ORDER_TABLE} ORDER BY CREATION_DATE ASC FETCH FIRST 1 ROW ONLY" # Get the oldest order
    result, response = exe_sql_query(ora_curx, sql_query)
    if not result or not response:
        return DEFAULT_ORDER_DATE
    else:
        date_str = (response[0][0] + dt.timedelta(days=1)).strftime("%d-%b-%y")
        logger.info (f"order date: {date_str}")
        return date_str


def generate_dynamic_fake_value(ora_curx, document_type, placeholder):
    """
    Generate dynamic fake values based on the placeholder name.
    :param placeholder: The string inside the '<>' to determine the value type.
    :return: A generated fake value or a default fake word if not recognized.
    """
    ord_num = 0
    ord_date = None
    
    # logger.info(f"placeholder: '{placeholder}'"
    if "gen_from_order" in placeholder:
        logger.info(f"placeholder: {placeholder}")
        pick_date = get_ord_date_from_table(ora_curx)
        return pick_date
    if "get_order_number" in placeholder:
        logger.info(f"placeholder: {placeholder}")
        ord_num = get_ord_po_from_table(ora_curx, document_type, 'ORDER')
        return ord_num
    if "get_po_number" in placeholder:
        logger.info(f"placeholder: {placeholder}")
        po_num = get_ord_po_from_table(ora_curx, document_type, 'PO')
        return po_num
    if "order number" in placeholder:
        logger.info(f"placeholder: {placeholder}")
        ord_num, ord_date = get_order_info_from_table(ora_curx)
        return ord_num
    elif "order date" in placeholder:
        # logger.info(f"placeholder: {placeholder} ----------->")
        if ord_num == 0:
            ord_num, ord_date = get_order_info_from_table(ora_curx)
        return ord_date
    elif "company" in placeholder:
        return faker.company()
    elif "address" in placeholder:
        return faker.address().replace("\n", ", ")
    elif "first name" in placeholder:
        return faker.first_name()
    elif "last name" in placeholder:
        return faker.last_name()
    elif "name" in placeholder:
        return faker.name()
    elif "phone" in placeholder:
        return faker.phone_number()
    elif "email" in placeholder:
        return faker.email()
    elif "date" in placeholder:
        return datetime.now().strftime("%m-%d-%Y")
    elif "time" in placeholder:
        return datetime.now().strftime("%H:%M:%S")
    elif "alphanumeric" in placeholder:
        return faker.bothify("??###")
    elif "number" in placeholder:
        return str(faker.random_int(min=100000, max=999999))
    elif "sentence" in placeholder:
        return faker.sentence(nb_words=5)
    elif "word" in placeholder:
        return faker.word()
    else:
        return faker.word()  # Default fallback if no match is found

def replace_placeholders_preserve_formatting(ora_curx, doc_typ, doc):
    """
    Replace placeholders in the given Word document while preserving formatting.

    :param doc: The Word document object (loaded via python-docx).
    """
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # logger.info(f"paragraph> run.text: {run.text}")
            if '<' in run.text and '>' in run.text:
                run.text = re.sub(
                    r"<(.*?)>",
                    lambda match: generate_dynamic_fake_value(ora_curx, doc_typ, match.group(1)),
                    run.text
                )

    # Replace placeholders in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        logger.info(f"table> run.text: {run.text}")
                        if '<' in run.text and '>' in run.text:
                            run.text = re.sub(
                                r"<(.*?)>",
                                lambda match: generate_dynamic_fake_value(ora_curx, doc_typ, match.group(1)),
                                run.text
                            )

def replace_placeholders_in_word(ora_curr, doctype, template_path, output_word_path, output_pdf_path):
    """
    Replace placeholders in a Word document and generate both Word and PDF outputs.

    :param template_path: Path to the input Word template.
    :param output_word_path: Path to save the modified Word document.
    :param output_pdf_path: Path to save the generated PDF file.
    """

    logger.info(f"output_pdf_path: {output_pdf_path}")
    # Load the Word document
    doc = Document(template_path)

    # Replace placeholders while preserving formatting
    replace_placeholders_preserve_formatting(ora_curr, doctype, doc)

    # Save the modified Word document
    doc.save(output_word_path)
    logger.info(f"Modified Word document saved to: {output_word_path}")

    # Convert Word to PDF
    try:
        import pypandoc
        pypandoc.convert_file(output_word_path, 'pdf', outputfile=output_pdf_path)
    except ImportError:
        from docx2pdf import convert
        convert(output_word_path, output_pdf_path)

    logger.info(f"Generated PDF saved to: {output_pdf_path}")

if __name__ == "__main__":

    orcle_conx = connect_to_oracle()
    orcle_curr = orcle_conx.cursor()
    result, response = exe_sql_query(orcle_curr, "SELECT * from V$VERSION")

    if not orcle_conx:
        logger.error("Oracle connection failed. So considering default values for order number and date ***")

    parser = argparse.ArgumentParser(description="Replace placeholders in a Word template and generate a PDF.")
    parser.add_argument("template", help="Path to the input Word template file.")
    parser.add_argument("output_word", help="Path to save the modified Word document.")
    parser.add_argument("output_pdf", help="Path to save the generated PDF file.")
    args = parser.parse_args()

    # Replace placeholders and generate files
    replace_placeholders_in_word(orcle_curr, doc_type, args.template, args.output_word, args.output_pdf)
    orcle_curr.close()
    orcle_conx.close()

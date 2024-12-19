import argparse
from docx import Document
from faker import Faker
import re
from datetime import datetime

# Initialize Faker for generating fake data
faker = Faker()

def generate_dynamic_fake_value(placeholder):
    """
    Generate dynamic fake values based on the placeholder name.
    :param placeholder: The string inside the '<>' to determine the value type.
    :return: A generated fake value or a default fake word if not recognized.
    """
    placeholder = placeholder.lower().strip()
    if "company" in placeholder:
        return faker.company()
    elif "address" in placeholder:
        return faker.address().replace("\n", ", ")
    elif "first name" in placeholder:
        return faker.first_name()
    elif "last name" in placeholder:
        return faker.last_name()
    elif "name" in placeholder:
        return faker.name()
    elif "phone" in placeholder:
        return faker.phone_number()
    elif "email" in placeholder:
        return faker.email()
    elif "date" in placeholder:
        return datetime.now().strftime("%d-%m-%Y")
    elif "time" in placeholder:
        return datetime.now().strftime("%H:%M:%S")
    elif "alphanumeric" in placeholder:
        return faker.bothify("??###")
    elif "number" in placeholder:
        return str(faker.random_int(min=100000, max=999999))
    elif "sentence" in placeholder:
        return faker.sentence(nb_words=5)
    elif "word" in placeholder:
        return faker.word()
    else:
        return faker.word()  # Default fallback if no match is found

def replace_placeholders_preserve_formatting(doc):
    """
    Replace placeholders in the given Word document while preserving formatting.

    :param doc: The Word document object (loaded via python-docx).
    """
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if '<' in run.text and '>' in run.text:
                run.text = re.sub(
                    r"<(.*?)>",
                    lambda match: generate_dynamic_fake_value(match.group(1)),
                    run.text
                )

    # Replace placeholders in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if '<' in run.text and '>' in run.text:
                            run.text = re.sub(
                                r"<(.*?)>",
                                lambda match: generate_dynamic_fake_value(match.group(1)),
                                run.text
                            )

def replace_placeholders_in_word(template_path, output_word_path, output_pdf_path):
    """
    Replace placeholders in a Word document and generate both Word and PDF outputs.

    :param template_path: Path to the input Word template.
    :param output_word_path: Path to save the modified Word document.
    :param output_pdf_path: Path to save the generated PDF file.
    """
    # Load the Word document
    doc = Document(template_path)

    # Replace placeholders while preserving formatting
    replace_placeholders_preserve_formatting(doc)

    # Save the modified Word document
    doc.save(output_word_path)
    print(f"Modified Word document saved to: {output_word_path}")

    # Convert Word to PDF
    try:
        import pypandoc
        pypandoc.convert_file(output_word_path, 'pdf', outputfile=output_pdf_path)
    except ImportError:
        from docx2pdf import convert
        convert(output_word_path, output_pdf_path)

    print(f"Generated PDF saved to: {output_pdf_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Replace placeholders in a Word template and generate a PDF.")
    parser.add_argument("template", help="Path to the input Word template file.")
    parser.add_argument("output_word", help="Path to save the modified Word document.")
    parser.add_argument("output_pdf", help="Path to save the generated PDF file.")
    args = parser.parse_args()

    # Replace placeholders and generate files
    replace_placeholders_in_word(args.template, args.output_word, args.output_pdf)
